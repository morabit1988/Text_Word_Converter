import os
import tkinter as tk
from tkinter import filedialog, messagebox

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from fpdf import FPDF

from license_checker import is_pro_user


# === INTERFACE ===

root = tk.Tk()
root.title("Convertisseur TXT → DOCX / PDF")
root.geometry("450x300")


# === CHOIX DU FICHIER TXT ===

def choose_file():
    file = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt")])
    if file:
        entry_file.delete(0, tk.END)
        entry_file.insert(0, file)


label_file = tk.Label(root, text="Fichier TXT à convertir :")
label_file.pack(pady=5)

entry_file = tk.Entry(root, width=50)
entry_file.pack()

button_browse = tk.Button(root, text="Parcourir", command=choose_file)
button_browse.pack(pady=5)


# === NOM DU FICHIER DE SORTIE ===

label_output = tk.Label(root, text="Nom du fichier de sortie (sans extension) :")
label_output.pack(pady=5)

entry_output = tk.Entry(root, width=50)
entry_output.pack()


# === OPTIONS DE FORMAT ===

var_docx = tk.BooleanVar()
var_pdf = tk.BooleanVar()

check_docx = tk.Checkbutton(root, text="Générer Word (.docx)", variable=var_docx)
check_pdf = tk.Checkbutton(root, text="Générer PDF (.pdf)", variable=var_pdf)

check_docx.pack()
check_pdf.pack()


# === FONCTIONS DE CONVERSION ===

def add_header(section, text="Document Header"):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = text
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(12)
    run.font.name = 'Arial'


def add_footer_with_page_numbers(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def convert_txt_to_word(input_file, output_file):
    try:
        with open(input_file, 'r', encoding='utf-8') as file:
            content = file.read()
    except FileNotFoundError:
        messagebox.showerror("Erreur", f"Fichier introuvable : {input_file}")
        return

    pro = is_pro_user()
    if not pro:
        content = content[:1000]

    doc = Document()
    section = doc.sections[0]
    add_header(section, text="Conversion TXT → DOCX")
    add_footer_with_page_numbers(section)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if (stripped.isupper() or stripped.startswith('#') or
                stripped.lower().startswith("titre") or stripped.endswith(":")):
            doc.add_paragraph(stripped, style='Heading 1')
        else:
            doc.add_paragraph(stripped, style='Normal')

    doc.save(output_file)


def txt_to_pdf(input_file, output_file):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)

    try:
        with open(input_file, 'r', encoding='utf-8') as file:
            content = file.read()
    except FileNotFoundError:
        messagebox.showerror("Erreur", f"Fichier introuvable : {input_file}")
        return

    paragraphs = content.split('\n\n')
    for para in paragraphs:
        cleaned = para.strip().replace('\n', ' ')
        if cleaned:
            pdf.multi_cell(0, 10, cleaned)
            pdf.ln(5)

    pdf.output(output_file)


# === ACTION : CONVERTIR ===

def convert():
    input_file = entry_file.get()
    output_name = entry_output.get().strip()

    if not input_file or not output_name:
        messagebox.showwarning("Champs manquants", "Veuillez sélectionner un fichier et entrer un nom de sortie.")
        return

    if not var_docx.get() and not var_pdf.get():
        messagebox.showwarning("Format manquant", "Veuillez cocher au moins un format (DOCX ou PDF).")
        return

    if var_docx.get():
        convert_txt_to_word(input_file, output_name + ".docx")

    if var_pdf.get():
        txt_to_pdf(input_file, output_name + ".pdf")

    messagebox.showinfo("Succès", "Conversion terminée avec succès ✅")


button_convert = tk.Button(root, text="Convertir", command=convert, bg="lightblue")
button_convert.pack(pady=15)


# === LANCEMENT DE L’INTERFACE ===

root.mainloop()
